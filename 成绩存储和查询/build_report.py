from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
OUT.mkdir(exist_ok=True)
REPORT = OUT / "基于HBase的成绩存储和查询系统课程大作业报告.docx"
SHOTS = Path(r"C:\腾讯电脑管家截图文件")
BLUE, GRAY, LIGHT = "1F4D78", "666666", "EAF1F8"

def font(run, name="宋体", size=11, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("ascii", "hAnsi", "eastAsia", "cs"): rf.set(qn("w:"+a), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def para(doc, text="", size=11, align=None, indent=True, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if indent: p.paragraph_format.first_line_indent = Cm(.74)
    if align is not None: p.alignment = align
    font(p.add_run(text), "宋体", size)
    return p

def heading(doc, text, lvl=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if lvl == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if lvl == 1 else 5)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), "黑体", 16 if lvl == 1 else 13, True, BLUE)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(.74); p.paragraph_format.right_indent = Cm(.74)
    p.paragraph_format.space_after = Pt(7); p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text); font(r, "Consolas", 8.5)
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F3F5F7"); p._p.get_or_add_pPr().append(shd)
    return p

def caption(doc, text):
    p = para(doc, text, 9.5, WD_ALIGN_PARAGRAPH.CENTER, False, 9)
    for r in p.runs: font(r, "宋体", 9.5, color=GRAY)

def figure(doc, file, text, width=6.1):
    path = SHOTS / file
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
        p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, text)

def set_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    grid = table._tbl.tblGrid
    for col, w in zip(grid.gridCol_lst, widths): col.set(qn("w:w"), str(w))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w/1440)
            tcw = cell._tc.tcPr.tcW; tcw.set(qn("w:w"), str(w)); tcw.set(qn("w:type"), "dxa")
            tcpr = cell._tc.get_or_add_tcPr()
            mar = OxmlElement("w:tcMar")
            for s in ("top","start","bottom","end"):
                el = OxmlElement("w:"+s); el.set(qn("w:w"), "100" if s in ("top","bottom") else "120"); el.set(qn("w:type"), "dxa"); mar.append(el)
            tcpr.append(mar)

def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; set_widths(t, widths)
    for c, h in zip(t.rows[0].cells, headers):
        shd=OxmlElement("w:shd"); shd.set(qn("w:fill"), LIGHT); c._tc.get_or_add_tcPr().append(shd)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; font(c.paragraphs[0].add_run(h), "宋体", 10, True)
    for row in rows:
        cells=t.add_row().cells
        for c, v in zip(cells,row): font(c.paragraphs[0].add_run(v), "宋体", 10)
    return t

def footer(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("《大数据分析与处理基础实践》课程大作业报告"), "宋体", 9, color=GRAY)

doc=Document()
sec=doc.sections[0]
sec.page_width,sec.page_height=Cm(21),Cm(29.7)
sec.top_margin=sec.bottom_margin=Cm(2.54); sec.left_margin=sec.right_margin=Cm(2.54)
sec.header_distance=sec.footer_distance=Cm(1.25); footer(sec)
normal=doc.styles["Normal"]; normal.font.name="宋体"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); normal.font.size=Pt(11)

# cover
for _ in range(5): doc.add_paragraph()
p=para(doc,"重庆交通大学信息科学与工程学院",18,WD_ALIGN_PARAGRAPH.CENTER,False,24); font(p.runs[0],"黑体",18,True)
p=para(doc,"课 程 大 作 业",28,WD_ALIGN_PARAGRAPH.CENTER,False,46); font(p.runs[0],"黑体",28,True,BLUE)
p=para(doc,"基于 HBase 的成绩存储和查询系统设计与实现",20,WD_ALIGN_PARAGRAPH.CENTER,False,42); font(p.runs[0],"黑体",20,True)
for label,value in [("课程名称","大数据分析与处理基础实践"),("专业班级","________________"),("学    号","________________"),("姓    名","________________"),("任课教师","________________")]:
    p=para(doc,label+"："+value,14,WD_ALIGN_PARAGRAPH.CENTER,False,10); font(p.runs[0],"宋体",14)
for _ in range(3): doc.add_paragraph()
p=para(doc,"2026 年 6 月",14,WD_ALIGN_PARAGRAPH.CENTER,False); font(p.runs[0],"宋体",14)
doc.add_page_break()

# contents
p=para(doc,"目 录",22,WD_ALIGN_PARAGRAPH.CENTER,False,20); font(p.runs[0],"黑体",22,True)
for s in ["第一章  概述","  1.1 任务概述","  1.2 数据集描述","第二章  生成模拟数据集","  2.1 目标与设计","  2.2 生成方法与关键代码","第三章  功能实现","  3.1 HBase 表结构与行键设计","  3.2 成绩数据导入","  3.3 按学号和课程号查询","第四章  运行及测试","  4.1 环境配置与启动验证","  4.2 编译、打包与数据导入","  4.3 功能测试与结果分析","第五章  总结"]:
    p=para(doc,s,12,indent=False,after=5); p.paragraph_format.left_indent=Cm(1); font(p.runs[0],"宋体",12)
doc.add_page_break()

heading(doc,"第一章  概述")
heading(doc,"1.1 任务概述",2)
para(doc,"本课程大作业以“成绩存储和查询”为题，使用 Java 和 HBase 实现模拟成绩数据的生成、列族数据库存储，以及按照学号和课程号进行查询。系统在单机 HBase 2.4.17 环境中完成从数据生成、表结构设计、批量导入到 HBase Shell 和 Java 程序验证的完整流程。")
para(doc,"项目的目标是掌握 HBase 的表、列族、行键、Put、Scan 和 PrefixFilter 等基本概念，并用真实运行结果验证程序的正确性。")
heading(doc,"1.2 数据集描述",2)
para(doc,"数据文件为 data/scores.txt，每行由“学号 课程号 成绩”三个字段构成。学号按 2026 加 6 位序号生成，课程号范围为 C001 至 C016，成绩范围为 30 至 100。生成器固定随机种子 20260622，以保证数据可复现。")
table(doc,["字段","说明","示例"],[("studentId","学生学号","2026000001"),("courseId","课程号","C003"),("score","成绩","30-100")],[1800,4500,3060]); caption(doc,"表 1  模拟成绩数据集字段说明")
para(doc,"生成器有意限制学号和课程号的随机范围，使部分“学号+课程号”组合重复，从而模拟补考和重修场景。导入程序通过记录序号保存这些多次成绩，不会覆盖已有数据。")

heading(doc,"第二章  生成模拟数据集")
heading(doc,"2.1 目标与设计",2)
para(doc,"模拟数据集不少于 10000 条记录，同时应具有随机性、格式统一性和重复组合。程序随机产生 1200 名学生、16 门课程和 30-100 分的成绩，最终生成 10000 条 UTF-8 编码记录。")
heading(doc,"2.2 生成方法与关键代码",2)
para(doc,"ScoreDataGenerator 在每次循环中生成学号、课程号和成绩，并以空格分隔写入文本文件。程序开始时检查记录数必须不少于 10000 条。关键逻辑如下：")
code(doc,'String studentId = String.format(Locale.ROOT, "2026%06d", 1 + random.nextInt(1_200));\nString courseId = COURSES[random.nextInt(COURSES.length)];\nint score = 30 + random.nextInt(71);\nwriter.write(studentId + " " + courseId + " " + score);')
heading(doc,"2.3 数据质量分析",2)
para(doc,"数据覆盖多个学生和课程组合，成绩取值合法。重复组合用于模拟实际业务中的重修或补考；每一行格式统一，便于导入程序使用空白符或逗号分隔解析。")

heading(doc,"第三章  功能实现")
heading(doc,"3.1 HBase 表结构与行键设计",2)
para(doc,"系统使用表名 cs，列族为 info。行键设计为“studentId#courseId#attempt”，例如 2026000001#C003#00001。attempt 为同一学生同一课程的第几次记录，固定宽度序号既避免覆盖，又便于按学号前缀扫描。")
table(doc,["列","用途"],[("info:studentId","保存学号"),("info:courseId","保存课程号"),("info:score","保存成绩"),("info:attempt","保存第几次记录"),("info:importedAt","保存导入时间")],[3000,6360]); caption(doc,"表 2  cs 表的列族与列限定符设计")
heading(doc,"3.2 成绩数据导入",2)
para(doc,"导入功能逐行读取 scores.txt，为每一个“学号#课程号”维护出现次数，构造不重复的行键；随后通过 Put 对象写入全部列。程序先检测 cs 表是否存在，仅在不存在时创建列族 info，避免重复建表。")
heading(doc,"3.3 按学号和课程号查询",2)
para(doc,"按学号查询使用 PrefixFilter，前缀为“学号#”，可筛选一个学生的全部课程成绩。按课程号查询使用 Scan 遍历结果后比较 info:courseId 的实际值，输出匹配课程的学生成绩。前者直接受益于行键设计，后者满足功能需求，并可在后续通过二级索引进一步优化。")

heading(doc,"第四章  运行及测试")
heading(doc,"4.1 环境配置与启动验证",2)
para(doc,"实验在 VMware 虚拟机 HBase-work 中完成，使用 Ubuntu 18.04、OpenJDK 1.8.0_362 和 HBase 2.4.17 单机模式。HBase 数据根目录设置为 /home/chenh/hbase-data，内嵌 ZooKeeper 数据目录设置为 /home/chenh/zookeeper-data。为适配本地文件系统，配置中关闭 hflush 能力的严格检查。")
code(doc,"<property>\n  <name>hbase.unsafe.stream.capability.enforce</name>\n  <value>false</value>\n</property>")
para(doc,"执行 start-hbase.sh 后，在 HBase Shell 中运行 status 'simple'。输出显示 active master、1 live servers 和 0 dead servers，说明服务已正常启动。")
figure(doc,"局部截取_20260622_185632.png","图 1  HBase 单机服务状态验证（1 个 live server）")
heading(doc,"4.2 编译、打包与数据导入",2)
para(doc,"在项目目录使用 Maven 执行 mvn clean package，生成 score-storage-query-1.0.0.jar。随后使用 HBase 的 RunJar 方式运行导入程序，程序自动创建 cs 表并成功写入 10000 条成绩记录。")
code(doc,"hbase org.apache.hadoop.util.RunJar target/score-storage-query-1.0.0.jar import data/scores.txt")
figure(doc,"局部截取_20260622_191950.png","图 2  HBase Shell 中确认 cs 表已创建")
figure(doc,"局部截取_20260622_192045.png","图 3  cs 表列族 info 的结构验证")
figure(doc,"局部截取_20260622_192056.png","图 4  count 命令验证共写入 10000 条记录")
heading(doc,"4.3 功能测试与结果分析",2)
para(doc,"使用 scan 'cs', { LIMIT => 10 } 抽样查看行键及 info 列族内容，可见 attempt、courseId、importedAt、score 和 studentId 等字段均正确保存。使用 ROWPREFIXFILTER 查询学号 2026000001，可获得该学生的多个课程成绩。")
figure(doc,"局部截取_20260622_192119.png","图 5  scan 抽样验证行键和 info 列族数据")
figure(doc,"局部截取_20260622_192151.png","图 6  按学号前缀筛选的 HBase Shell 验证")
para(doc,"最后运行 Java 查询命令。学号 2026000001 返回 7 条成绩记录；课程 C001 查询输出多个学生的成绩及重修次数。结果表明表设计、数据导入和两种查询功能均正确实现。")
figure(doc,"局部截取_20260622_192225.png","图 7  Java 程序按学号查询结果")
figure(doc,"局部截取_20260622_192240.png","图 8  Java 程序按课程号查询结果（节选）")

heading(doc,"4.4 环境配置问题与解决",2)
para(doc,"为完整记录项目实施过程，本节汇总前期旧虚拟机的配置问题，以及在最终 HBase-work 环境中遇到的启动和打包兼容性问题。前期问题说明了配置文件和运行路径必须保持一致；最终运行结果以第 4.1 至 4.3 节的成功验证为准。")
table(doc,["问题现象","原因定位","处理措施与结果"],[
    ("虚拟机无法联网：apt update 报 Could not resolve us.archive.ubuntu.com","虚拟机 DNS 或网络配置不正确，无法访问软件源。","改用 Windows 主机下载所需安装包，再通过 VMware 文件共享或复制粘贴导入 Ubuntu，完成离线部署。"),
    ("HBase 目录提示 No such file or directory","HBASE_HOME 与压缩包的实际解压层级不一致，conf、bin 路径无法定位。","检查解压目录后统一 HBASE_HOME；最终环境采用 /home/chenh/hbase。"),
    ("HBase 启动提示 /usr/local/bin/java: No such file or directory","JAVA_HOME 指向了不存在的 Java 可执行文件路径。","通过 java -version 和实际 JDK 目录确认，最终设置为 /usr/lib/jvm/java-8-openjdk-amd64。"),
    ("hbase-site.xml 提示 Illegal to have multiple roots 或不合法注释","编辑 XML 时保留了旧内容，造成多个根节点或不符合 XML 规范的注释。","完整替换为仅含一个 configuration 根节点的合法配置文件。"),
    ("status 'simple' 报 ConnectionLoss for /hbase/master","旧环境中 ZooKeeper 2181 端口或残留进程发生冲突；最终环境日志还发现 WAL 写入的 hflush 能力校验失败，导致 HMaster 主动退出。","旧环境停止残留进程并统一配置；最终环境设置 hbase.unsafe.stream.capability.enforce=false 后，服务显示 1 live servers。"),
    ("Maven 打包依次提示编译器/ Shade 插件要求 Maven 3.6.3，Java 8 不支持 --release","Ubuntu 18.04 自带 Maven 较旧，而初始插件版本和 release 参数偏新。","将 compiler 插件调整为 3.8.1、shade 插件调整为 3.2.4，并使用 source/target 1.8 编译，最终 BUILD SUCCESS。")
],[2200,3100,4060]); caption(doc,"表 3  环境配置与排错过程汇总")
para(doc,"日志中还出现了“Class path contains multiple SLF4J bindings”和“Unable to load native-hadoop library”等警告。经分析，这些属于日志组件重复绑定或本地库缺失的常见警告，HBase 会回退使用内置 Java 类；在服务状态、建表、导入和查询均正常的前提下，不影响本项目的核心功能。")
figure(doc,"局部截取_20260622_152457.png","图 9  排错阶段出现的 ConnectionLoss 提示（后续已解决）")
figure(doc,"局部截取_20260622_191323.png","图 10  Maven Shade 插件版本不兼容提示（后续已解决）")

heading(doc,"第五章  总结")
para(doc,"本项目完成了基于 HBase 的成绩存储和查询系统。通过 Java 程序生成 10000 条模拟成绩，采用“学号#课程号#记录序号”行键避免多次成绩覆盖，并以 info 列族保存成绩字段。实际运行表明 HBase 服务稳定、cs 表结构正确、10000 条数据成功导入，Shell 与 Java 查询均返回预期结果。")
para(doc,"实现过程中还完成了单机 HBase、ZooKeeper、WAL 写入约束以及 Maven/Java 兼容性配置与排查。这说明大数据处理项目不仅需要业务代码，也需要关注运行环境、数据模型和测试过程。未来可建立课程维度索引，降低按课程号查询的全表扫描成本。")

heading(doc,"附录  主要运行命令")
code(doc,"start-hbase.sh\nhbase shell\nstatus 'simple'\nlist\ndescribe 'cs'\ncount 'cs'\nscan 'cs', { LIMIT => 10 }\nhbase org.apache.hadoop.util.RunJar target/score-storage-query-1.0.0.jar student 2026000001\nhbase org.apache.hadoop.util.RunJar target/score-storage-query-1.0.0.jar course C001")
doc.core_properties.title="基于HBase的成绩存储和查询系统设计与实现"
doc.core_properties.author="chenh"
doc.save(REPORT)
print(REPORT)
